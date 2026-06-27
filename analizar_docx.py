from docx import Document
doc = Document(r'C:\Users\jhvar\Downloads\formulario 2\LISTA DE CHEQUEO INGENIERÍA CIVIL 1.docx')
print('=== PARAGRAFOS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print('[{}] style={} | {}'.format(i, p.style.name, repr(p.text)))
print()
print('=== TABLAS ===')
for ti, table in enumerate(doc.tables):
    print('--- Tabla {} ({} filas x {} cols) ---'.format(ti, len(table.rows), len(table.columns)))
    for ri, row in enumerate(table.rows):
        row_texts = []
        for cell in row.cells:
            row_texts.append(repr(cell.text.strip()[:80]))
        print('  Fila {}: {}'.format(ri, ' | '.join(row_texts)))
