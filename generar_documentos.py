"""
Generador de documentos IDENTICOS a los originales.
Usa 2.docx, 3.xlsx como plantillas y rellena con datos del candidato.
Para el form 4 genera un .docx siguiendo la estructura del 4.pdf.
"""

import os, re, copy, shutil
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SS_ID          = "1fXU5t9fmDfXwskFs42r1eZNZa0KCxNo1Li77yrDpyvY"
CARPETA_SALIDA = "Documentos_Candidatos"
PLANTILLA_2    = "2.docx"
PLANTILLA_3    = "3.xlsx"

def autenticar():
    creds = Credentials.from_authorized_user_file("token.json",
        ["https://www.googleapis.com/auth/spreadsheets"])
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
    return build("sheets", "v4", credentials=creds)

def leer_hoja(service, nombre_hoja):
    res = service.spreadsheets().values().get(
        spreadsheetId=SS_ID, range=nombre_hoja + "!A1:ZZ500").execute()
    filas = res.get("values", [])
    if len(filas) < 2:
        return []
    enc = filas[0]
    return [dict(zip(enc, f + [""] * (len(enc) - len(f)))) for f in filas[1:]]

def safe(d, k, default=""):
    return str(d.get(k, default) or default).strip()

def nombre_archivo(nombre, cedula):
    n = re.sub(r"[^\w\s-]", "", nombre).strip().replace(" ", "_")
    return f"{cedula}_{n}"

# ═══════════════════════════════════════════════════════
# FORM 2 → .docx  usando 2.docx como plantilla EXACTA
# ═══════════════════════════════════════════════════════
def generar_docx_form2(fila, carpeta):
    cedula   = safe(fila, "Cedula del Candidato")
    nombre   = safe(fila, "Nombre Completo del Candidato")
    programa = safe(fila, "Programa / Area del Concurso")
    concepto = safe(fila, "Concepto Final")
    obs_gral = safe(fila, "Observaciones Generales de la Verificacion")

    # Mapeo de requisitos del formulario a filas de la tabla del docx
    requisitos_mapeados = [
        safe(fila, "1. Formato de Inscripcion firmado por el aspirante"),
        safe(fila, "2. Hoja de Vida UQ (Formato GH-FOR-006)"),
        safe(fila, "3. Fotocopia de cedula y libreta militar"),
        safe(fila, "4. Fotocopia de matricula o tarjeta profesional"),
        safe(fila, "5. Certificados disciplinarios, judiciales o fiscales vigentes"),
        safe(fila, "6. Certificado de experiencia docente universitaria (minimo 2 anos / 1024 horas catedra)"),
        safe(fila, "7. Titulo de Pregrado requerido por el perfil"),
        safe(fila, "8. Titulo de Posgrado requerido por el perfil"),
        safe(fila, "9. Experiencia minima en el area del concurso (segun perfil)"),
        safe(fila, "10. Tema de disertacion presentado (si aplica)"),
    ]
    obs_mapeadas = [
        safe(fila, "Observaciones - Formato de Inscripcion"),
        safe(fila, "Observaciones - Hoja de Vida UQ"),
        safe(fila, "Observaciones - Cedula / Libreta Militar"),
        safe(fila, "Observaciones - Matricula / Tarjeta Profesional"),
        safe(fila, "Observaciones - Certificados disciplinarios"),
        safe(fila, "Observaciones - Experiencia Docente"),
        safe(fila, "Observaciones - Titulo Pregrado"),
        safe(fila, "Observaciones - Titulo Posgrado"),
        safe(fila, "Observaciones - Experiencia en el Area"),
        safe(fila, "Observaciones - Tema de Disertacion"),
    ]

    # Abrir plantilla original
    doc = Document(PLANTILLA_2)

    # Rellenar parrafos de cabecera
    for para in doc.paragraphs:
        if "NOMBRE:" in para.text and "C.C." in para.text:
            # Limpiar runs y reescribir
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"NOMBRE: {nombre}    C.C. {cedula}"
        elif "FACULTAD DE" in para.text.upper() or "FACULTAD:" in para.text.upper():
            for run in para.runs:
                run.text = ""
            if para.runs:
                # Extraer facultad del programa (primer elemento antes del guion)
                facultad = programa.split(" - ")[0] if " - " in programa else programa
                para.runs[0].text = f"FACULTAD: {facultad}"
        elif "PROGRAMA:" in para.text.upper():
            for run in para.runs:
                run.text = ""
            if para.runs:
                prog_nombre = programa.split(" - ")[-1] if " - " in programa else programa
                para.runs[0].text = f"PROGRAMA: {prog_nombre}"
        elif "AREA O LINEA:" in para.text.upper():
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"AREA O LINEA: {safe(fila, 'Perfil del Cargo')}"
        elif "OBSERVACIONES GENERALES:" in para.text.upper():
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"OBSERVACIONES GENERALES: {obs_gral}"

    # Rellenar Tabla 0 (requisitos) — filas 1 a 10
    tabla_req = doc.tables[0]
    for i, (estado, obs) in enumerate(zip(requisitos_mapeados, obs_mapeadas)):
        if i + 1 < len(tabla_req.rows):
            fila_tabla = tabla_req.rows[i + 1]
            # Columna OBSERVACIONES (1)
            celda_obs = fila_tabla.cells[1]
            for run in celda_obs.paragraphs[0].runs:
                run.text = ""
            if celda_obs.paragraphs[0].runs:
                celda_obs.paragraphs[0].runs[0].text = obs
            else:
                celda_obs.paragraphs[0].add_run(obs)
            # Columna CUMPLE (2)
            celda_cumple = fila_tabla.cells[2]
            for run in celda_cumple.paragraphs[0].runs:
                run.text = ""
            cumple_val = "SI" if "CUMPLE" in estado.upper() and estado.upper()[:3] != "NO " else "NO"
            if celda_cumple.paragraphs[0].runs:
                celda_cumple.paragraphs[0].runs[0].text = cumple_val
            else:
                celda_cumple.paragraphs[0].add_run(cumple_val)

    # Rellenar Tabla 1 (concepto final SI/NO)
    tabla_conc = doc.tables[1]
    if len(tabla_conc.rows) > 1:
        fila_conc = tabla_conc.rows[1]
        si_no = "SI" if "CUMPLE CON TODOS" in concepto.upper() else "NO"
        if si_no == "SI":
            fila_conc.cells[1].paragraphs[0].clear()
            fila_conc.cells[1].paragraphs[0].add_run("X")
        else:
            fila_conc.cells[2].paragraphs[0].clear()
            fila_conc.cells[2].paragraphs[0].add_run("X")

    nombre_f = os.path.join(carpeta, nombre_archivo(nombre, cedula) + "_ETAPA2_Verificacion.docx")
    doc.save(nombre_f)
    print(f"  DOCX Etapa 2: {os.path.basename(nombre_f)}")


# ═══════════════════════════════════════════════════════
# FORM 3 → .xlsx  usando 3.xlsx como plantilla EXACTA
# ═══════════════════════════════════════════════════════
def generar_xlsx_form3(fila, carpeta):
    cedula   = safe(fila, "Cedula del Candidato")
    nombre   = safe(fila, "Nombre Completo del Candidato")
    programa = safe(fila, "Programa / Area del Concurso")
    evaluador = safe(fila, "Nombre del Evaluador (Miembro del Comite)")

    pts_c1  = safe(fila, "Puntaje Total Criterio 1 - Nivel Academico (0 a 12 pts)")
    pts_c2  = safe(fila, "Puntaje Total Criterio 2 - Experiencia (0 a 15 pts)")
    pts_c3  = safe(fila, "Puntaje Total Criterio 3 - Productividad (0 a 8 pts)")

    nivel_acad  = safe(fila, "Nivel Academico Acreditado")
    just_acad   = safe(fila, "Justificacion - Nivel Academico")
    exp_doc     = safe(fila, "2a. Experiencia Docente Universitaria (hasta 5 puntos)")
    exp_inv     = safe(fila, "2b. Experiencia en Investigacion (hasta 4 puntos)")
    exp_ext     = safe(fila, "2c. Experiencia en Extension o Proyeccion Social (hasta 2 puntos)")
    exp_prof    = safe(fila, "2d. Experiencia Profesional Diferente a la Docente (hasta 2 puntos)")
    exp_cargos  = safe(fila, "2e. Experiencia en Cargos Academico Administrativos (hasta 2 puntos)")
    det_arts    = safe(fila, "Detalle de articulos indexados")
    det_libros  = safe(fila, "Detalle de libros / obras / software")
    obs         = safe(fila, "Observaciones Generales del Evaluador")

    def extraer_pts(s):
        for tok in s.replace("->","").replace("→","").split():
            tok_c = tok.replace(",", ".")
            try:
                return float(tok_c)
            except:
                pass
        return ""

    # Cargar plantilla original
    wb = openpyxl.load_workbook(PLANTILLA_3)
    ws = wb.active

    def escribir_celda(ws, coord, valor):
        """Escribe en la celda sin romper el formato existente."""
        try:
            ws[coord] = valor
        except:
            pass

    # Rellenar encabezado (filas 2-8 según estructura real del 3.xlsx)
    # Fila con Nombre
    for row in ws.iter_rows():
        for cell in row:
            v = str(cell.value or "")
            if "Nombre:" in v and len(v) < 60:
                cell.value = f"Nombre:  {nombre}"
            elif "Facultad:" in v and "INGENIERÍA" in v:
                fac = programa.split(" - ")[0] if " - " in programa else programa
                cell.value = f"Facultad: {fac}"
            elif "Programa:" in v and "INGENIERIA CIVIL" in v:
                prog = programa.split(" - ")[-1] if " - " in programa else programa
                cell.value = f"Programa: {prog}"
            elif "Área o linea:" in v or "Area o linea:" in v:
                cell.value = f"Área o linea: {evaluador}"

    # Escribir el puntaje de Nivel Académico
    # En el xlsx original la columna E contiene el puntaje vacío
    # Buscar la celda donde va el Nivel Académico acreditado y el puntaje
    # Según estructura: fila 9 col C = nivel, fila 9 col D = puntaje criterio
    for row in ws.iter_rows():
        for cell in row:
            v = str(cell.value or "")
            # Nivel académico (título)
            if "Maestría Ingeniería" in v or "Doctorado Ingeniería" in v:
                # Celda de puntaje está a la derecha
                pass  # La plantilla ya tiene los criterios, se llenan a mano

    # Llenar el resumen de puntajes (filas 69-73)
    for row in ws.iter_rows():
        for cell in row:
            v = str(cell.value or "")
            if v == "1" and str(ws.cell(row=cell.row, column=cell.column+2).value or "") == "Máximo nivel Académico obtenido":
                ws.cell(row=cell.row, column=cell.column+3).value = pts_c1 or extraer_pts(nivel_acad)
            if v == "2" and str(ws.cell(row=cell.row, column=cell.column+2).value or "") == "Experiencia":
                ws.cell(row=cell.row, column=cell.column+3).value = pts_c2
            if v == "3" and str(ws.cell(row=cell.row, column=cell.column+2).value or "") == "Productividad Académica":
                ws.cell(row=cell.row, column=cell.column+3).value = pts_c3
            if "Total hoja de vida" in v:
                try:
                    total = (float(str(pts_c1).replace(",",".") or 0) +
                             float(str(pts_c2).replace(",",".") or 0) +
                             float(str(pts_c3).replace(",",".") or 0))
                    ws.cell(row=cell.row, column=cell.column+2).value = total
                except:
                    pass

    # Agregar hoja de detalle con las respuestas textuales del evaluador
    if "Detalle" not in wb.sheetnames:
        ws2 = wb.create_sheet("Detalle Evaluacion")
    else:
        ws2 = wb["Detalle Evaluacion"]
    ws2["A1"] = "DETALLE DE LA EVALUACION"
    ws2["A1"].font = Font(bold=True, size=12)
    ws2["A2"] = f"Candidato: {nombre}    CC: {cedula}"
    ws2["A3"] = f"Evaluador: {evaluador}"
    ws2["A4"] = f"Programa: {programa}"
    ws2["A6"]  = "NIVEL ACADEMICO SELECCIONADO:"
    ws2["B6"]  = nivel_acad
    ws2["A7"]  = "Justificacion nivel:"
    ws2["B7"]  = just_acad
    ws2["A9"]  = "2a. Experiencia Docente:"
    ws2["B9"]  = exp_doc
    ws2["A10"] = "2b. Investigacion:"
    ws2["B10"] = exp_inv
    ws2["A11"] = "2c. Extension:"
    ws2["B11"] = exp_ext
    ws2["A12"] = "2d. Exp. Profesional:"
    ws2["B12"] = exp_prof
    ws2["A13"] = "2e. Cargos Academicos:"
    ws2["B13"] = exp_cargos
    ws2["A15"] = "Articulos indexados:"
    ws2["B15"] = det_arts
    ws2["A16"] = "Libros / obras:"
    ws2["B16"] = det_libros
    ws2["A18"] = "Observaciones generales:"
    ws2["B18"] = obs
    ws2["A20"] = f"PUNTAJE C1 (Nivel Academico):   {pts_c1} / 12"
    ws2["A21"] = f"PUNTAJE C2 (Experiencia):        {pts_c2} / 15"
    ws2["A22"] = f"PUNTAJE C3 (Productividad):      {pts_c3} / 8"
    try:
        total = (float(str(pts_c1).replace(",",".") or 0) +
                 float(str(pts_c2).replace(",",".") or 0) +
                 float(str(pts_c3).replace(",",".") or 0))
        ws2["A23"] = f"TOTAL HOJA DE VIDA:              {total} / 35"
        ws2["A23"].font = Font(bold=True, size=11)
    except:
        pass
    for col in ["A","B"]:
        ws2.column_dimensions[col].width = 35

    nombre_f = os.path.join(carpeta, nombre_archivo(nombre, cedula) + "_ETAPA3_Calificacion.xlsx")
    wb.save(nombre_f)
    print(f"  XLSX Etapa 3: {os.path.basename(nombre_f)}")


# ═══════════════════════════════════════════════════════
# FORM 4 → .docx  estructura identica a 4.pdf
# ═══════════════════════════════════════════════════════
def generar_docx_form4(fila, carpeta):
    nombre    = safe(fila, "Nombre Completo")
    cedula    = safe(fila, "Cedula de Ciudadania")
    programa  = safe(fila, "Programa Academico")
    facultad  = safe(fila, "Facultad")
    anno      = safe(fila, "Anno del Concurso")
    categoria = safe(fila, "Categoria de Ingreso").upper()
    pregrado  = safe(fila, "Titulo de Pregrado")
    inst_pre  = safe(fila, "Institucion Pregrado")
    pts_pre   = safe(fila, "Puntaje Titulo Pregrado")
    posgrado  = safe(fila, "Titulo de Posgrado (Maestria o Doctorado)")
    inst_pos  = safe(fila, "Institucion Posgrado")
    pts_pos   = safe(fila, "Puntaje Titulo Posgrado")
    sub_tit   = safe(fila, "SUBTOTAL TITULOS (puntos)")
    pts_cat   = safe(fila, "Puntos por Categoria")
    inv       = safe(fila, "3.1 Investigacion (puntos)")
    det_inv   = safe(fila, "Detalle Investigacion")
    doc_uni   = safe(fila, "3.2 Docencia Universitaria (puntos)")
    det_doc   = safe(fila, "Detalle Docencia Universitaria")
    cargos    = safe(fila, "3.3 Experiencia en Cargos de Direccion Academica (puntos)")
    det_carg  = safe(fila, "Detalle Cargos Direccion")
    exp_prof  = safe(fila, "3.4 Experiencia Profesional (puntos)")
    det_prof  = safe(fila, "Detalle Experiencia Profesional")
    sub_exp   = safe(fila, "SUBTOTAL EXPERIENCIA CALIFICADA (puntos)")
    arts      = safe(fila, "Articulos en Revistas Indexadas")
    libros    = safe(fila, "Libros y Capitulos de Libro")
    sub_prod  = safe(fila, "SUBTOTAL PRODUCTIVIDAD ACADEMICA (puntos)")
    total_pts = safe(fila, "TOTAL PUNTOS FINALES")
    remun     = safe(fila, "Remuneracion Mensual (pesos colombianos)")
    proyecto  = safe(fila, "Proyecto (nombre del funcionario que proyecta la ficha)")
    aprobo    = safe(fila, "Aprobo (nombre del Jefe de Asuntos Profesorales)")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def add_bold_center(text, size=11):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        return p

    def add_table_row(table, col0, col1, col2="", bold0=False):
        row = table.add_row()
        row.cells[0].text = col0
        row.cells[1].text = col1
        if col2 != "" and len(row.cells) > 2:
            row.cells[2].text = col2
        if bold0 and row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        return row

    # ── Encabezado ────────────────────────────────────────
    add_bold_center(f"FICHA DE INGRESO A LA CARRERA DOCENTE - AÑO {anno}", 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"NOMBRE: {nombre.upper()}   C.C. {cedula}").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"PROGRAMA: {programa.upper()}   FACULTAD: {facultad.upper()}").bold = True

    cat_texto = f"{categoria}={pts_cat} PTS.   TOPE EXP. CALIFICADA= {sub_exp} PTS"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(cat_texto)

    doc.add_paragraph()

    # ── Tabla principal (estructura idéntica al PDF) ───────
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(10)
    table.columns[2].width = Cm(3)

    def add_row(c0, c1, c2, bold=False):
        row = table.add_row()
        row.cells[0].text = c0
        row.cells[1].text = c1
        row.cells[2].text = str(c2)
        if bold:
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                else:
                    cell.paragraphs[0].add_run(cell.text).bold = True
                    cell.text = ""
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return row

    # Titulos
    add_row("", "PUNTOS", "", bold=True)
    add_row("PREGRADO", f"{inst_pre}   {pregrado}", pts_pre)
    add_row("1. TÍTULOS", "", "")
    add_row("POSTGRADO", f"{inst_pos}   {posgrado}", pts_pos)

    # Categoria
    add_row("2. CATEGORÍA", f"{categoria}  Cumple con los requisitos para la categoria", pts_cat)

    # Experiencia
    add_row("3.1 INVESTIGACIÓN", det_inv, inv)
    add_row("3.2 DOCENCIA UNIVERSITARIA", det_doc, doc_uni)
    add_row("3. EXPERIENCIA CALIFICADA  3.3 CARGOS DIRECCIÓN ACADÉMICA",
            det_carg if det_carg else "N/P", cargos if cargos else "")
    add_row("3.4 EXPERIENCIA PROFESIONAL", det_prof if det_prof else "N/P", exp_prof if exp_prof else "")
    add_row("SUBTOTAL", "", sub_exp, bold=True)

    # Productividad — cada articulo en una fila
    arts_lista = [a.strip() for a in arts.split("\n") if a.strip()] if arts else ["N/P"]
    for i, art in enumerate(arts_lista):
        label = "4. PRODUCTIVIDAD ACADÉMICA  ARTÍCULOS" if i == 0 else ""
        add_row(label, art, "")

    libros_lista = [l.strip() for l in libros.split("\n") if l.strip()] if libros else []
    for lib in libros_lista:
        add_row("LIBROS", lib, "")

    add_row("SUBTOTAL", "", sub_prod, bold=True)

    doc.add_paragraph()

    # ── Totales ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.add_run(f"TOTAL PUNTOS: {total_pts}").bold = True
    p.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # Texto narrativo (como en el PDF)
    try:
        total_f = float(str(total_pts).replace(",","."))
        p_narr = doc.add_paragraph()
        p_narr.add_run(
            f"Con {pts_pre} puntos por título de Pregrado, {pts_pos} puntos por el título de Posgrado, "
            f"{sub_prod} puntos por productividad académica, {pts_cat} puntos por la categoría de "
            f"profesor {categoria.title()} y {sub_exp} puntos por experiencia calificada, "
            f"para un total de {total_pts} puntos y una remuneración mensual de:"
        )
        p_rem = doc.add_paragraph()
        p_rem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rem.add_run(f"$ {remun} MONEDA CORRIENTE").bold = True
        p_rem.runs[0].font.size = Pt(11)
    except:
        p = doc.add_paragraph(f"Remuneración mensual: $ {remun}")

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Firmas ─────────────────────────────────────────────
    t_firma = doc.add_table(rows=3, cols=2)
    t_firma.style = "Table Grid"
    t_firma.rows[0].cells[0].text = proyecto
    t_firma.rows[0].cells[1].text = aprobo
    t_firma.rows[1].cells[0].text = "PROYECTÓ:"
    t_firma.rows[1].cells[1].text = "APROBÓ:"
    t_firma.rows[2].cells[0].text = "Contratista Oficina de Asuntos Profesorales"
    t_firma.rows[2].cells[1].text = "Jefe Oficina Asuntos Profesorales"
    for row in t_firma.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    nombre_f = os.path.join(carpeta, nombre_archivo(nombre, cedula) + "_ETAPA4_FichaIngreso.docx")
    doc.save(nombre_f)
    print(f"  DOCX Etapa 4: {os.path.basename(nombre_f)}")


# ═══════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════
def main():
    print("=" * 65)
    print("  GENERADOR DE DOCUMENTOS — ESTRUCTURA IDENTICA A ORIGINALES")
    print("=" * 65)

    service = autenticar()

    datos_f2 = leer_hoja(service, "Respuestas de formulario 2")
    datos_f3 = leer_hoja(service, "Respuestas de formulario 3")
    datos_f4 = leer_hoja(service, "Respuestas de formulario 4")

    print(f"\n  Respuestas Etapa 2: {len(datos_f2)}")
    print(f"  Respuestas Etapa 3: {len(datos_f3)}")
    print(f"  Respuestas Etapa 4: {len(datos_f4)}")

    os.makedirs(CARPETA_SALIDA, exist_ok=True)
    total = 0

    print(f"\n  Generando en: {os.path.abspath(CARPETA_SALIDA)}")
    print("-" * 65)

    for fila in datos_f2:
        nombre = safe(fila, "Nombre Completo del Candidato")
        if nombre:
            carpeta_c = os.path.join(CARPETA_SALIDA,
                nombre_archivo(nombre, safe(fila, "Cedula del Candidato")))
            os.makedirs(carpeta_c, exist_ok=True)
            generar_docx_form2(fila, carpeta_c)
            total += 1

    for fila in datos_f3:
        nombre = safe(fila, "Nombre Completo del Candidato")
        if nombre:
            carpeta_c = os.path.join(CARPETA_SALIDA,
                nombre_archivo(nombre, safe(fila, "Cedula del Candidato")))
            os.makedirs(carpeta_c, exist_ok=True)
            generar_xlsx_form3(fila, carpeta_c)
            total += 1

    for fila in datos_f4:
        nombre = safe(fila, "Nombre Completo")
        if nombre:
            carpeta_c = os.path.join(CARPETA_SALIDA,
                nombre_archivo(nombre, safe(fila, "Cedula de Ciudadania")))
            os.makedirs(carpeta_c, exist_ok=True)
            generar_docx_form4(fila, carpeta_c)
            total += 1

    print("-" * 65)
    print(f"\n  Total documentos generados: {total}")
    print(f"  Carpeta: {os.path.abspath(CARPETA_SALIDA)}")
    print("=" * 65)

    import subprocess
    subprocess.Popen(f'explorer "{os.path.abspath(CARPETA_SALIDA)}"')

if __name__ == "__main__":
    main()
