from docx import Document
from docx.oxml.ns import qn
import zipfile, json

def rgb_from_hex(h):
    if not h or h == 'auto': return None
    return '#' + h.upper()

doc = Document(r'C:\Users\jhvar\Downloads\formulario 2\LISTA DE CHEQUEO INGENIERÍA CIVIL 1.docx')

print("=== PARRAFOS (texto completo) ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print("[{}] '{}'".format(i, p.text.strip()))

print()
print("=== ESTILOS DE TABLA ===")
for ti, table in enumerate(doc.tables):
    print("--- TABLA {} ({} filas x {} cols) ---".format(ti, len(table.rows), len(table.columns)))
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()[:50]
            # Obtener color de fondo de la celda
            tc = cell._tc
            shd = tc.find('.//' + qn('w:shd'))
            bg = None
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                bg = rgb_from_hex(fill) if fill else None
            # Obtener alineacion del primer parrafo
            aln = None
            if cell.paragraphs:
                p = cell.paragraphs[0]
                if p.alignment:
                    aln = str(p.alignment)
            # Obtener negrita del primer run
            bold = False
            if cell.paragraphs and cell.paragraphs[0].runs:
                bold = cell.paragraphs[0].runs[0].bold
            print("  [{},{}] bg={} bold={} aln={} | '{}'".format(ri, ci, bg, bold, aln, txt))

print()
print("=== PROPIEDADES TABLA 0 (widths) ===")
table0 = doc.tables[0]
for ci in range(len(table0.columns)):
    col = table0.columns[ci]
    print("  Col {}: width={}".format(ci, col.width))
